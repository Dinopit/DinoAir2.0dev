"""
Text extraction module for DinoAir 2.0 RAG File Search system.
Provides modular text extractors for various file formats.
"""

import os
import json
import csv
import chardet
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List
from pathlib import Path

# Import logging from DinoAir's logger
from ..utils.logger import Logger


class BaseExtractor(ABC):
    """
    Abstract base class for all text extractors.
    Defines the interface that all extractors must implement.
    """
    
    # Default max file size: 50MB
    DEFAULT_MAX_FILE_SIZE = 50 * 1024 * 1024
    
    def __init__(self, max_file_size: Optional[int] = None):
        """
        Initialize the base extractor.
        
        Args:
            max_file_size: Maximum file size to process in bytes.
                          Defaults to 50MB.
        """
        self.logger = Logger()
        self.max_file_size = max_file_size or self.DEFAULT_MAX_FILE_SIZE
        self.supported_extensions = []
    
    def can_extract(self, file_path: str) -> bool:
        """
        Check if this extractor can handle the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if this extractor can handle the file
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions
    
    def check_file_size(self, file_path: str) -> bool:
        """
        Check if file size is within acceptable limits.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if file size is acceptable
        """
        try:
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                self.logger.warning(
                    f"File {file_path} exceeds size limit: "
                    f"{file_size} > {self.max_file_size}"
                )
                return False
            return True
        except Exception as e:
            self.logger.error(f"Error checking file size: {str(e)}")
            return False
    
    @abstractmethod
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from the given file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict containing:
                - success: bool
                - text: extracted text (if successful)
                - metadata: additional metadata (if available)
                - error: error message (if failed)
        """
        pass
    
    def detect_encoding(self, file_path: str) -> str:
        """
        Detect file encoding using chardet.
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Detected encoding or 'utf-8' as fallback
        """
        try:
            with open(file_path, 'rb') as f:
                # Read first 10KB for detection
                raw_data = f.read(10240)
                if raw_data:
                    result = chardet.detect(raw_data)
                    encoding = result.get('encoding', 'utf-8')
                    confidence = result.get('confidence', 0)
                    
                    if confidence > 0.7:
                        return encoding
            
            return 'utf-8'
        except Exception as e:
            self.logger.error(f"Error detecting encoding: {str(e)}")
            return 'utf-8'


class TextFileExtractor(BaseExtractor):
    """
    Extractor for plain text files (.txt).
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.txt']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a plain text file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            return {
                "success": True,
                "text": text,
                "metadata": {
                    "encoding": encoding,
                    "file_type": "text"
                }
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract text: {str(e)}"
            }


class PDFExtractor(BaseExtractor):
    """
    Extractor for PDF files using PyPDF2.
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.pdf']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a PDF file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Import PyPDF2 only when needed
            try:
                import PyPDF2
            except ImportError:
                return {
                    "success": False,
                    "error": "PyPDF2 not installed. Please install it with: pip install PyPDF2"
                }
            
            text_content = []
            metadata = {
                "file_type": "pdf",
                "pages": []
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                metadata["total_pages"] = num_pages
                
                for page_num in range(num_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        
                        if page_text:
                            text_content.append(page_text)
                            metadata["pages"].append({
                                "page_number": page_num + 1,
                                "text_length": len(page_text)
                            })
                    except Exception as e:
                        self.logger.warning(
                            f"Error extracting page {page_num + 1}: {str(e)}"
                        )
                        metadata["pages"].append({
                            "page_number": page_num + 1,
                            "error": str(e)
                        })
            
            combined_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": combined_text,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract PDF: {str(e)}"
            }


class DocxExtractor(BaseExtractor):
    """
    Extractor for Microsoft Word documents (.docx) using python-docx.
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.docx']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a DOCX file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Import python-docx only when needed
            try:
                import docx
            except ImportError:
                return {
                    "success": False,
                    "error": "python-docx not installed. Please install it with: pip install python-docx"
                }
            
            doc = docx.Document(file_path)
            text_content = []
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            combined_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": combined_text,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract DOCX: {str(e)}"
            }


class CodeFileExtractor(BaseExtractor):
    """
    Extractor for common code files (.py, .js, .java, .cpp, .cs, etc.).
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = [
            '.py', '.js', '.java', '.cpp', '.c', '.cs', '.rb', '.go',
            '.php', '.swift', '.kt', '.rs', '.ts', '.jsx', '.tsx',
            '.vue', '.sh', '.bash', '.ps1', '.r', '.scala', '.h',
            '.hpp', '.cc', '.m', '.mm', '.pl', '.lua', '.dart'
        ]
        
        # Map extensions to language names for metadata
        self.language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.cs': 'csharp',
            '.rb': 'ruby',
            '.go': 'go',
            '.php': 'php',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.rs': 'rust',
            '.ts': 'typescript',
            '.jsx': 'javascript-react',
            '.tsx': 'typescript-react',
            '.vue': 'vue',
            '.sh': 'shell',
            '.bash': 'bash',
            '.ps1': 'powershell',
            '.r': 'r',
            '.scala': 'scala',
            '.h': 'c-header',
            '.hpp': 'cpp-header',
            '.cc': 'cpp',
            '.m': 'objective-c',
            '.mm': 'objective-cpp',
            '.pl': 'perl',
            '.lua': 'lua',
            '.dart': 'dart'
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a code file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            # Get file extension and language
            ext = Path(file_path).suffix.lower()
            language = self.language_map.get(ext, 'unknown')
            
            # Count lines and detect if it has comments
            lines = text.split('\n')
            line_count = len(lines)
            
            metadata = {
                "encoding": encoding,
                "file_type": "code",
                "language": language,
                "line_count": line_count,
                "extension": ext
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting code from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract code: {str(e)}"
            }


class MarkdownExtractor(BaseExtractor):
    """
    Extractor for Markdown files (.md).
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.md', '.markdown']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a Markdown file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            # Extract headers for metadata
            headers = []
            for line in text.split('\n'):
                if line.strip().startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    header_text = line.strip('#').strip()
                    if header_text:
                        headers.append({
                            "level": level,
                            "text": header_text
                        })
            
            metadata = {
                "encoding": encoding,
                "file_type": "markdown",
                "headers": headers,
                "header_count": len(headers)
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting markdown from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract markdown: {str(e)}"
            }


class JSONExtractor(BaseExtractor):
    """
    Extractor for JSON files.
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.json']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a JSON file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            # Read and parse JSON
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                json_data = json.load(f)
            
            # Convert JSON to formatted text
            text = json.dumps(json_data, indent=2, ensure_ascii=False)
            
            # Extract metadata about the JSON structure
            metadata = {
                "encoding": encoding,
                "file_type": "json",
                "root_type": type(json_data).__name__,
                "key_count": len(json_data) if isinstance(json_data, dict) else None,
                "array_length": len(json_data) if isinstance(json_data, list) else None
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata
            }
            
        except json.JSONDecodeError as e:
            self.logger.error(f"Invalid JSON in {file_path}: {str(e)}")
            # Still try to read as plain text
            try:
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    text = f.read()
                return {
                    "success": True,
                    "text": text,
                    "metadata": {
                        "encoding": encoding,
                        "file_type": "json",
                        "parse_error": str(e)
                    }
                }
            except Exception as e2:
                return {
                    "success": False,
                    "error": f"Failed to extract JSON: {str(e2)}"
                }
        except Exception as e:
            self.logger.error(f"Error extracting JSON from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract JSON: {str(e)}"
            }


class CSVExtractor(BaseExtractor):
    """
    Extractor for CSV files.
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        super().__init__(max_file_size)
        self.supported_extensions = ['.csv']
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a CSV file."""
        try:
            if not self.check_file_size(file_path):
                return {
                    "success": False,
                    "error": "File exceeds size limit"
                }
            
            # Detect encoding
            encoding = self.detect_encoding(file_path)
            
            text_lines = []
            row_count = 0
            column_count = 0
            headers = []
            
            # Read CSV file
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                # Detect delimiter
                sample = f.read(1024)
                f.seek(0)
                
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample)
                    delimiter = dialect.delimiter
                except:
                    delimiter = ','
                
                csv_reader = csv.reader(f, delimiter=delimiter)
                
                for i, row in enumerate(csv_reader):
                    if i == 0:
                        headers = row
                        column_count = len(row)
                    
                    # Convert row to text
                    row_text = delimiter.join(row)
                    text_lines.append(row_text)
                    row_count += 1
            
            text = '\n'.join(text_lines)
            
            metadata = {
                "encoding": encoding,
                "file_type": "csv",
                "delimiter": delimiter,
                "row_count": row_count,
                "column_count": column_count,
                "headers": headers
            }
            
            return {
                "success": True,
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting CSV from {file_path}: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract CSV: {str(e)}"
            }


class ExtractorFactory:
    """
    Factory class to create appropriate extractors based on file type.
    """
    
    def __init__(self, max_file_size: Optional[int] = None):
        """
        Initialize the factory with all available extractors.
        
        Args:
            max_file_size: Maximum file size for all extractors
        """
        self.logger = Logger()
        self.extractors = [
            TextFileExtractor(max_file_size),
            PDFExtractor(max_file_size),
            DocxExtractor(max_file_size),
            CodeFileExtractor(max_file_size),
            MarkdownExtractor(max_file_size),
            JSONExtractor(max_file_size),
            CSVExtractor(max_file_size)
        ]
        
        # Build extension mapping for quick lookup
        self.extension_map = {}
        for extractor in self.extractors:
            for ext in extractor.supported_extensions:
                self.extension_map[ext] = extractor
    
    def get_extractor(self, file_path: str) -> Optional[BaseExtractor]:
        """
        Get the appropriate extractor for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            BaseExtractor instance or None if no suitable extractor found
        """
        ext = Path(file_path).suffix.lower()
        extractor = self.extension_map.get(ext)
        
        if not extractor:
            self.logger.warning(f"No extractor found for extension: {ext}")
        
        return extractor
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if file type is supported
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.extension_map
    
    def get_supported_extensions(self) -> List[str]:
        """
        Get list of all supported file extensions.
        
        Returns:
            List of supported extensions
        """
        return sorted(list(self.extension_map.keys()))